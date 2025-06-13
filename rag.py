
import requests
import json

import os
from openai import OpenAI
import environ



from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.retrievers import BM25Retriever
import os
from langchain.docstore.document import Document
from docx import Document as DocxDocument
from pypdf import PdfReader

# 使用本地文件夹知识库。路径为../knowledge
import os

# 加载本地 doc/pdf 并转为 Document 对象
def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return full_text

def load_pdf(file_path):
    reader = PdfReader(file_path)
    full_text = ""
    for page in reader.pages:
        full_text += page.extract_text() or ""
    return full_text

def load_documents_from_folder(folder_path):
    docs = []
    for fname in os.listdir(folder_path):
        fpath = os.path.join(folder_path, fname)
        if fname.lower().endswith(".docx"):
            text = load_docx(fpath)
        elif fname.lower().endswith(".pdf"):
            text = load_pdf(fpath)
        else:
            continue  # 跳过非 docx/pdf 文件
        docs.append(
            Document(
                page_content=text,
                metadata={"source": fname}
            )
        )
    return docs

# 使用你的本地知识库路径
import os

# 获取 main.py 所在目录
base_dir = os.path.dirname(os.path.abspath(__file__))

# 计算 knowledge_base 文件夹的路径（上级目录的 knowledge_base 文件夹）
folder_path = os.path.join(base_dir, ".", "knowledge")

# 标准化路径（防止 ../.. 出现问题）
folder_path = os.path.normpath(folder_path)

print(folder_path)
# 输出类似于 /home/youruser/my_agent_project/knowledge_base

source_docs = load_documents_from_folder(folder_path)
print(f"Loaded {len(source_docs)} base documents")

# Split documents into smaller chunks for better retrieval
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,  # Characters per chunk
    chunk_overlap=50,  # Overlap between chunks to maintain context
    add_start_index=True,
    strip_whitespace=True,
    separators=["\n\n", "\n", ".", " ", ""],  # Priority order for splitting
)
docs_processed = text_splitter.split_documents(source_docs)

print(f"Knowledge base prepared with {len(docs_processed)} document chunks")

from smolagents import Tool

class RetrieverTool(Tool):
    name = "retriever"
    description = "Uses semantic search to retrieve the parts of transformers documentation that could be most relevant to answer your query."
    inputs = {
        "query": {
            "type": "string",
            "description": "The query to perform. This should be semantically close to your target documents. Use the affirmative form rather than a question.",
        }
    }
    output_type = "string"

    def __init__(self, docs, **kwargs):
        super().__init__(**kwargs)
        # Initialize the retriever with our processed documents
        self.retriever = BM25Retriever.from_documents(
            docs, k=10  # Return top 10 most relevant documents
        )

    def forward(self, query: str) -> str:
        """Execute the retrieval based on the provided query."""
        assert isinstance(query, str), "Your search query must be a string"

        # Retrieve relevant documents
        docs = self.retriever.invoke(query)

        # Format the retrieved documents for readability
        return "\nRetrieved documents:\n" + "".join(
            [
                f"\n\n===== Document {str(i)} =====\n" + doc.page_content
                for i, doc in enumerate(docs)
            ]
        )

# Initialize our retriever tool with the processed documents
retriever_tool = RetrieverTool(docs_processed)

from smolagents import InferenceClientModel, CodeAgent
import os
from smolagents import OpenAIServerModel

model = OpenAIServerModel(
    model_id="qwen-plus",  # Use the DashScope model name
    api_base=os.environ["OPENAI_API_BASE"] ,
    api_key=os.environ["OPENAI_API_KEY"],
)
agent = CodeAgent(
    tools=[retriever_tool],  # List of tools available to the agent
    model=model,
    max_steps=4,  # Limit the number of reasoning steps
    verbosity_level=2,  # Show detailed agent reasoning
)

# Ask a question that requires retrieving information
# question = "我要出差去省内，公杂费是每天多少?"
question = "我要出差去北京，我是省部级人员，住宿费标准是每天多少?"

# Run the agent to get an answer
agent_output = agent.run(question)

# Display the final answer
print("\nFinal answer:")
print(agent_output)
